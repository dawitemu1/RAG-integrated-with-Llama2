import os
import streamlit as st
import torch
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
import docx  # python-docx
import numpy as np

from langchain.chains import RetrievalQA
from langchain_community.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain_community.llms import Ollama  # << NEW

# -------------------------- TESSERACT FIX --------------------------
pytesseract.pytesseract.tesseract_cmd = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
# ------------------------------------------------------------------

os.environ["TRANSFORMERS_NO_TF"] = "1"
FAISS_DB_PATH = "faiss_db"

# ✅ Use Ollama with llama2 model
llm = Ollama(model="llama2")

# Load MiniLM embedding model for FAISS
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Function to extract text from uploaded files
def extract_text(file):
    file_type = file.name.lower()
    try:
        if file_type.endswith(".pdf"):
            text = ""
            doc = fitz.open(stream=file.read(), filetype="pdf")
            for page in doc:
                text += page.get_text()
            return text
        elif file_type.endswith(".docx"):
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        elif file_type.endswith((".png", ".jpg", ".jpeg")):
            image = Image.open(file)
            return pytesseract.image_to_string(image)
        else:
            return None
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return None

# Sidebar file upload
st.sidebar.title("\U0001F4E4 Upload Files here")
uploaded_files = st.sidebar.file_uploader(
    "Upload PDF, DOCX, or Image files",
    type=["pdf", "docx", "png", "jpg", "jpeg"],
    accept_multiple_files=True
)

if uploaded_files:
    new_docs = []
    for file in uploaded_files:
        extracted_text = extract_text(file)
        if extracted_text:
            new_docs.append(Document(page_content=extracted_text, metadata={"source": file.name}))
        else:
            st.sidebar.warning(f"Couldn't extract text from {file.name}")

    if new_docs:
        st.sidebar.success(f"{len(new_docs)} documents processed and added to FAISS!")
        texts = [doc.page_content for doc in new_docs]
        metadatas = [doc.metadata for doc in new_docs]

        if os.path.exists(FAISS_DB_PATH):
            vector_store = FAISS.load_local(FAISS_DB_PATH, embedding_model, allow_dangerous_deserialization=True)
            vector_store.add_texts(texts, metadatas)
        else:
            vector_store = FAISS.from_texts(texts, embedding_model, metadatas=metadatas)

        vector_store.save_local(FAISS_DB_PATH)

# Load FAISS vector store
if os.path.exists(FAISS_DB_PATH):
    try:
        vector_store = FAISS.load_local(FAISS_DB_PATH, embedding_model, allow_dangerous_deserialization=True)
        retriever = vector_store.as_retriever()
    except Exception as e:
        st.error(f"Error loading FAISS database: {e}")
        st.stop()
else:
    st.error("FAISS database not found. Please upload and process documents first.")
    st.stop()

# Create QA chain
qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)

# Streamlit Chat UI
st.title("AI assistant for CBE Policy, Procedure & Process Guidance")
st.write("Ask questions about the Commercial Bank of Ethiopia's procedures and guidelines.")

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "question_count" not in st.session_state:
    st.session_state.question_count = 0

# Display chat history
for msg in st.session_state.chat_history:
    with st.chat_message("user"):
        st.markdown(msg["user"])
    with st.chat_message("assistant"):
        st.markdown(msg["bot"])

# Chat input
user_message = st.chat_input("Type your question and press Enter to Get Answer...")

if user_message:
    st.chat_message("user").markdown(user_message)

    prompt = (
        "You are a helpful assistant trained on Commercial Bank of Ethiopia's procedures and guidelines.\n"
        "Please provide concise and clear answers. Answer the following:\n\n"
        f"User: {user_message.strip()}\nAssistant:"
    )

    try:
        response = qa_chain.run(prompt)

        sources = []
        results = retriever.get_relevant_documents(user_message)
        if results:
            for result in results:
                source = result.metadata.get('source', 'Unknown document')
                sources.append(f"🔹 Source: {source}")
        else:
            sources.append("No relevant information found. Please modify your query or provide more context.")

        with st.chat_message("assistant"):
            st.write(f"\U0001F4A1 **Answer:** {response}")
            st.write("Document Retrieval Results:")
            for idx, source in enumerate(sources, 1):
                st.write(f"{idx}. {source}")

        st.session_state.chat_history.append({
            "user": user_message.strip(),
            "bot": f"\U0001F4A1 **Answer:** {response}\n" + "\n".join(sources)
        })

        if st.session_state.question_count == 0:
            st.info("I'm here to help more if you have additional questions!")

        st.session_state.question_count += 1

    except Exception as e:
        st.error(f"Error generating response: {e}")
